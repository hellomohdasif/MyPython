from docx import Document
import os
import random
import string
# # from datetime import date
# from datetime import datetime, timedelta
# from bs4 import BeautifulSoup
# from conData import topics
# import pdfkit
# from conData import topics
from docx2pdf import convert


from docx import Document
from datetime import datetime, timedelta

presentday = datetime.now()
nextday = presentday + timedelta(1)
d2 = presentday.strftime('%d-%m-%Y')
d3 = nextday.strftime('%d-%m-%Y')


k = 0
for k in range(0,1):

    def random_string_generator(str_size, allowed_chars):
        return ''.join(random.choice(allowed_chars) for x in range(str_size))


    chars = string.ascii_letters + string.ascii_uppercase + string.digits
    chars1 = string.digits
    chars2 = string.ascii_letters

    size1 = 4
    size2 = 5
    size3 = 8

    randomNum = (random_string_generator(size1, chars1))
    randomChar = (random_string_generator(size2, chars2))
    randomChar1 = (random_string_generator(size1, chars2))
    randomname = (random_string_generator(size3, chars2))


    onlineusers = (random_string_generator(size2, chars1))
    mega = (random_string_generator(size3, chars))


    # open the document
    # filename = "C:\\Users\\Asif\\Desktop\\1.docx"

    # doc = Document("C:\\Users\\Asif\\Desktop\\1.docx")
    # Dictionary = {"KOOL": randomname.upper(), "DATETIME":randomname.upper()}

    def replace_string(filename):
        doc = Document(filename)

        list = ['KOOL', 'DATETIME']
        list2 = [randomChar, d3]

        for p in doc.paragraphs:
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for j in range(0,len(inline)):
                    for i in range(0, len(list)):


                        inline[j].text = inline[j].text.replace(list[i], list2[i])


        doc.save('C:\\Users\\Asif\\Desktop\\dest1.docx')
        return 1
    replace_string("1.docx")


